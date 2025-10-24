"""Check actual content of competitor files"""
from docx import Document
import pdfplumber

print("=" * 80)
print("CINTOLO PDF CONTENT:")
print("=" * 80)
try:
    with pdfplumber.open('CINTOLO/CINTOLO 10-2025 Lista 31-1 - hasta 12.pdf') as pdf:
        print(f"Total pages: {len(pdf.pages)}")
        for i, page in enumerate(pdf.pages[:3]):  # First 3 pages
            print(f"\n--- PAGE {i+1} ---")
            text = page.extract_text()
            print(text[:1000] if text else "No text extracted")
            
            tables = page.extract_tables()
            if tables:
                print(f"\nFound {len(tables)} tables on page {i+1}")
                for j, table in enumerate(tables[:2]):  # First 2 tables
                    print(f"\nTable {j+1} preview:")
                    for row in table[:5]:  # First 5 rows
                        print(row)
except Exception as e:
    print(f"Error: {e}")

print("\n" + "=" * 80)
print("CINTOLO DOCX CONTENT:")
print("=" * 80)
try:
    doc = Document('CINTOLO/CINTOLO.docx')
    print("Paragraphs:")
    for i, para in enumerate(doc.paragraphs[:10]):
        print(f"{i}: {para.text}")
    
    print("\nTables:")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        for j, row in enumerate(table.rows[:5]):
            print([cell.text for cell in row.cells])
except Exception as e:
    print(f"Error: {e}")

print("\n" + "=" * 80)
print("ZALOZE DOCX CONTENT:")
print("=" * 80)
try:
    doc = Document('PRECIOS ZALOZE/ZALOZE.docx')
    print("Paragraphs:")
    for i, para in enumerate(doc.paragraphs[:10]):
        print(f"{i}: {para.text}")
    
    print("\nTables:")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        for j, row in enumerate(table.rows[:5]):
            print([cell.text for cell in row.cells])
except Exception as e:
    print(f"Error: {e}")


