"""
Extractor for Cintolo competitor price list.
Parses PDF and DOCX files from CINTOLO folder.
"""
import pandas as pd
import pdfplumber
from docx import Document
import logging
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class CintoloExtractor:
    def __init__(self, pdf_path='CINTOLO/CINTOLO 10-2025 Lista 31-1 - hasta 12.pdf',
                 docx_path='CINTOLO/CINTOLO.docx'):
        """Initialize with paths to Cintolo files."""
        self.pdf_path = pdf_path
        self.docx_path = docx_path
    
    def extract(self):
        """
        Extract competitor prices from Cintolo files.
        Returns a DataFrame with columns:
        - codigo: Product code
        - descripcion: Product description
        - precio: Cintolo price
        - tipo_serie: Product type/series
        - espesor: Thickness
        - size: Size
        """
        all_data = []
        
        # Try extracting from PDF
        try:
            pdf_data = self._extract_from_pdf()
            if not pdf_data.empty:
                all_data.append(pdf_data)
                logger.info(f"Extracted {len(pdf_data)} items from PDF")
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
        
        # Try extracting from DOCX
        try:
            docx_data = self._extract_from_docx()
            if not docx_data.empty:
                all_data.append(docx_data)
                logger.info(f"Extracted {len(docx_data)} items from DOCX")
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
        
        if not all_data:
            logger.warning("No data extracted from Cintolo files")
            return pd.DataFrame(columns=['codigo', 'descripcion', 'precio', 
                                        'tipo_serie', 'espesor', 'size'])
        
        # Merge all data
        merged_df = pd.concat(all_data, ignore_index=True)
        
        # Remove duplicates, keeping first occurrence
        if 'codigo' in merged_df.columns:
            merged_df = merged_df.drop_duplicates(subset=['codigo'], keep='first')
        
        logger.info(f"Total {len(merged_df)} unique products from Cintolo")
        return merged_df
    
    def _extract_from_pdf(self):
        """Extract data from PDF file."""
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                all_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        all_text.append(text)
                
                # Try to extract tables
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        all_tables.extend(tables)
                
                if all_tables:
                    return self._parse_pdf_tables(all_tables)
                else:
                    return self._parse_pdf_text('\n'.join(all_text))
                    
        except Exception as e:
            logger.error(f"Error in _extract_from_pdf: {e}")
            return pd.DataFrame()
    
    def _parse_pdf_tables(self, tables):
        """Parse tables extracted from PDF."""
        all_rows = []
        
        for table in tables:
            if not table or len(table) < 2:
                continue
            
            # First row is header: ARTIC., DESCRIPCION, ESPESOR, NORMA, CALIDAD, PRECIO USD
            headers = [str(h).strip().lower() if h else '' for h in table[0]]
            
            # Skip header row and process data rows
            for row in table[1:]:
                if not row or len(row) < 3:
                    continue
                
                try:
                    # Expected columns: ARTIC, DESCRIPCION, ESPESOR, NORMA, CALIDAD, PRECIO USD
                    artic = str(row[0]).strip() if row[0] else ''
                    descripcion = str(row[1]).strip() if len(row) > 1 and row[1] else ''
                    espesor = str(row[2]).strip() if len(row) > 2 and row[2] else ''
                    precio = str(row[5]).strip() if len(row) > 5 and row[5] else ''
                    
                    # Skip empty rows or invalid data
                    if not artic or not descripcion or not precio:
                        continue
                    
                    # Skip if it's a header repeat or invalid text
                    if artic.upper() == 'ARTIC.' or 'ARTIC' in artic.upper():
                        continue
                    if not artic.replace('.', '').isdigit():  # Article code should be numeric
                        continue
                    
                    # Convert price to float
                    try:
                        precio_float = float(precio.replace(',', '.'))
                        if precio_float <= 0:
                            continue
                    except:
                        continue
                    
                    # Extract size from description - handles dual sizes for reducers
                    import re
                    # First, try to extract dual sizes (e.g., "6" X 3"")
                    dual_size_match = re.search(r'(\d+(?:\.\d+)?(?:/\d+)?)\s*["\']?\s*[Xx×\*]\s*(\d+(?:\.\d+)?(?:/\d+)?)\s*["\']?', descripcion)
                    if dual_size_match:
                        size1 = dual_size_match.group(1).replace('.', ' ')
                        size2 = dual_size_match.group(2).replace('.', ' ')
                        size = f'{size1} X {size2}'
                    else:
                        # Single size (e.g., "2"", "1.1/2"")
                        size_match = re.search(r'(\d+(?:\.\d+)?(?:/\d+)?)\s*["\']', descripcion)
                        if not size_match:
                            # Try to match without quotes
                            size_match = re.search(r'\s(\d+(?:/\d+)?)\s*(?:"|$)', descripcion)
                        size = size_match.group(1).replace('.', ' ') if size_match else ''
                    
                    # Determine product type
                    desc_upper = descripcion.upper()
                    if 'RADIO CORTO 90' in desc_upper or 'CODO 90' in desc_upper:
                        tipo = 'CODOS 90°'
                    elif 'RADIO LARGO 90' in desc_upper:
                        tipo = 'CODOS 90° RADIO LARGO'
                    elif '45' in desc_upper:
                        tipo = 'CODOS 45°'
                    elif 'TE' in desc_upper or 'TEE' in desc_upper:
                        tipo = 'TE'
                    elif 'REDUCCION' in desc_upper or 'REDUCER' in desc_upper:
                        tipo = 'REDUCCION'
                    else:
                        tipo = desc_upper.split()[0] if desc_upper else 'OTROS'
                    
                    # Map espesor variations
                    espesor_clean = espesor.upper()
                    if 'STANDARD' in espesor_clean or 'STD' in espesor_clean:
                        espesor_clean = 'STD'
                    elif 'EXTRA PESADO' in espesor_clean or 'XS' in espesor_clean:
                        espesor_clean = 'XS'
                    elif 'LIVIANO' in espesor_clean:
                        espesor_clean = 'LIVIANO'
                    
                    all_rows.append({
                        'codigo': artic,
                        'descripcion': descripcion,
                        'precio': precio_float,
                        'tipo_serie': tipo,
                        'espesor': espesor_clean,
                        'size': size
                    })
                    
                except Exception as e:
                    logger.warning(f"Error processing row: {e}")
                    continue
        
        if not all_rows:
            return pd.DataFrame()
        
        df = pd.DataFrame(all_rows)
        return df
    
    def _parse_pdf_text(self, text):
        """Parse text extracted from PDF when tables are not available."""
        lines = text.split('\n')
        data = []
        
        for line in lines:
            # Look for patterns: code, description, price
            # This is a simplified parser - adjust based on actual format
            match = re.search(r'([A-Z0-9\-]+)\s+(.+?)\s+\$?\s*([\d,.]+)', line)
            if match:
                data.append({
                    'codigo': match.group(1),
                    'descripcion': match.group(2).strip(),
                    'precio': match.group(3).replace(',', '.')
                })
        
        if not data:
            return pd.DataFrame()
        
        df = pd.DataFrame(data)
        return self._normalize_cintolo_data(df)
    
    def _extract_from_docx(self):
        """Extract data from DOCX file."""
        try:
            doc = Document(self.docx_path)
            
            # Try extracting from tables first
            table_data = []
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    if i == 0:  # Skip header
                        continue
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2:
                        table_data.append(cells)
            
            if table_data:
                df = pd.DataFrame(table_data)
                df.columns = ['codigo', 'descripcion', 'precio'][:len(df.columns)]
                return self._normalize_cintolo_data(df)
            
            # If no tables, try parsing paragraphs
            para_data = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                match = re.search(r'([A-Z0-9\-]+)\s+(.+?)\s+\$?\s*([\d,.]+)', text)
                if match:
                    para_data.append({
                        'codigo': match.group(1),
                        'descripcion': match.group(2).strip(),
                        'precio': match.group(3).replace(',', '.')
                    })
            
            if para_data:
                df = pd.DataFrame(para_data)
                return self._normalize_cintolo_data(df)
            
            return pd.DataFrame()
            
        except Exception as e:
            logger.error(f"Error in _extract_from_docx: {e}")
            return pd.DataFrame()
    
    def _normalize_cintolo_data(self, df):
        """Normalize Cintolo data to standard format."""
        result_df = pd.DataFrame()
        
        # Find codigo column
        codigo_col = self._find_column(df, ['codigo', 'code', 'cod', 'item'])
        if codigo_col:
            result_df['codigo'] = df[codigo_col].astype(str)
        elif len(df.columns) > 0:
            result_df['codigo'] = df.iloc[:, 0].astype(str)
        
        # Find descripcion column
        desc_col = self._find_column(df, ['descripcion', 'description', 'desc'])
        if desc_col:
            result_df['descripcion'] = df[desc_col].astype(str)
        elif len(df.columns) > 1:
            result_df['descripcion'] = df.iloc[:, 1].astype(str)
        
        # Find precio column
        precio_col = self._find_column(df, ['precio', 'price', 'valor'])
        if precio_col:
            result_df['precio'] = pd.to_numeric(df[precio_col].astype(str).str.replace(',', '.'), 
                                               errors='coerce')
        elif len(df.columns) > 2:
            result_df['precio'] = pd.to_numeric(df.iloc[:, 2].astype(str).str.replace(',', '.'), 
                                               errors='coerce')
        
        # Extract tipo_serie, espesor, size from codigo or descripcion
        if 'codigo' in result_df.columns:
            result_df['tipo_serie'] = result_df['codigo'].apply(self._extract_tipo_serie)
        else:
            result_df['tipo_serie'] = ''
        
        if 'descripcion' in result_df.columns:
            result_df['espesor'] = result_df['descripcion'].apply(self._extract_espesor)
            result_df['size'] = result_df['descripcion'].apply(self._extract_size)
        else:
            result_df['espesor'] = ''
            result_df['size'] = ''
        
        # Remove invalid rows
        result_df = result_df.dropna(subset=['precio'])
        result_df = result_df[result_df['precio'] > 0]
        
        return result_df
    
    def _find_column(self, df, possible_names):
        """Find column by possible names."""
        df_cols = [str(c).strip().lower() for c in df.columns]
        for name in possible_names:
            if name in df_cols:
                idx = df_cols.index(name)
                return df.columns[idx]
        return None
    
    def _extract_tipo_serie(self, codigo):
        """Extract tipo_serie from product code."""
        try:
            parts = str(codigo).split('-')
            if len(parts) > 0:
                return parts[0].strip().upper()
        except:
            pass
        return ''
    
    def _extract_espesor(self, description):
        """Extract thickness from description."""
        try:
            desc = str(description).upper()
            match = re.search(r'SCH\s*(\d+|STD|XS|XXS)', desc)
            if match:
                return match.group(0).replace(' ', '')
        except:
            pass
        return ''
    
    def _extract_size(self, description):
        """Extract size from description - handles dual sizes for reducers."""
        try:
            desc = str(description)
            
            # First, try to extract dual sizes for reducers (e.g., "6" X 3"", "2.1/2" X 1"")
            # Match patterns like: 6" X 3", 2.1/2" X 1", 1/2" x 1/4", etc.
            dual_match = re.search(r'(\d+(?:\.\d+)?(?:/\d+)?)\s*["\']?\s*[Xx×\*]\s*(\d+(?:\.\d+)?(?:/\d+)?)\s*["\']?', desc)
            if dual_match:
                size1 = dual_match.group(1)
                size2 = dual_match.group(2)
                return f'{size1} X {size2}'
            
            # Single size with quotes (e.g., 6", 1/2")
            match = re.search(r'(\d+(?:\.\d+)?(?:/\d+)?)\s*["\']', desc)
            if match:
                return match.group(1)
            
            # Size in MM or INCH
            match = re.search(r'(\d+)\s*(MM|INCH)', desc, re.IGNORECASE)
            if match:
                return match.group(1) + match.group(2).upper()
        except:
            pass
        return ''

