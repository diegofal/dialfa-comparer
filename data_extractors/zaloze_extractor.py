"""
Extractor for Zaloze competitor price list.
Parses DOCX and JPEG files using EasyOCR for images.
"""
import pandas as pd
from docx import Document
import logging
import re

try:
    from PIL import Image
    import easyocr
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("EasyOCR libraries not available. Install easyocr and Pillow for image processing.")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ZalozeExtractor:
    def __init__(self, docx_path='PRECIOS ZALOZE/ZALOZE.docx',
                 jpeg_paths=None):
        """Initialize with paths to Zaloze files."""
        self.docx_path = docx_path
        if jpeg_paths is None:
            self.jpeg_paths = [
                'PRECIOS ZALOZE/WhatsApp Image 2025-10-22 at 12.38.37.jpeg',
                'PRECIOS ZALOZE/WhatsApp Image 2025-10-22 at 12.39.03.jpeg'
            ]
        else:
            self.jpeg_paths = jpeg_paths
    
    def extract(self):
        """
        Extract competitor prices from Zaloze files.
        Returns a DataFrame with columns:
        - codigo: Product code
        - descripcion: Product description
        - precio: Zaloze price
        - tipo_serie: Product type/series
        - espesor: Thickness
        - size: Size
        """
        all_data = []
        
        # Try extracting from DOCX
        try:
            docx_data = self._extract_from_docx()
            if not docx_data.empty:
                all_data.append(docx_data)
                logger.info(f"Extracted {len(docx_data)} items from DOCX")
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
        
        # Try extracting from JPEG images
        if OCR_AVAILABLE:
            for jpeg_path in self.jpeg_paths:
                try:
                    jpeg_data = self._extract_from_jpeg(jpeg_path)
                    if not jpeg_data.empty:
                        all_data.append(jpeg_data)
                        logger.info(f"Extracted {len(jpeg_data)} items from {jpeg_path}")
                except Exception as e:
                    logger.error(f"Error extracting from {jpeg_path}: {e}")
        else:
            logger.warning("Skipping JPEG extraction - OCR not available")
        
        if not all_data:
            logger.warning("No data extracted from Zaloze files")
            return pd.DataFrame(columns=['codigo', 'descripcion', 'precio', 
                                        'tipo_serie', 'espesor', 'size'])
        
        # Merge all data
        merged_df = pd.concat(all_data, ignore_index=True)
        
        # Remove duplicates
        if 'codigo' in merged_df.columns:
            merged_df = merged_df.drop_duplicates(subset=['codigo'], keep='first')
        
        logger.info(f"Total {len(merged_df)} unique products from Zaloze")
        return merged_df
    
    def _extract_from_docx(self):
        """Extract data from DOCX file."""
        try:
            doc = Document(self.docx_path)
            
            # Try extracting from tables
            table_data = []
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    if i == 0:  # Skip header
                        continue
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2:
                        table_data.append(cells)
            
            if table_data:
                df = pd.DataFrame(table_data)
                df.columns = ['codigo', 'descripcion', 'precio'][:len(df.columns)]
                return self._normalize_zaloze_data(df)
            
            # Try parsing paragraphs
            para_data = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                match = re.search(r'([A-Z0-9\-]+)\s+(.+?)\s+\$?\s*([\d,.]+)', text)
                if match:
                    para_data.append({
                        'codigo': match.group(1),
                        'descripcion': match.group(2).strip(),
                        'precio': match.group(3).replace(',', '.')
                    })
            
            if para_data:
                df = pd.DataFrame(para_data)
                return self._normalize_zaloze_data(df)
            
            return pd.DataFrame()
            
        except Exception as e:
            logger.error(f"Error in _extract_from_docx: {e}")
            return pd.DataFrame()
    
    def _extract_from_jpeg(self, jpeg_path):
        """Extract data from JPEG image using EasyOCR."""
        if not OCR_AVAILABLE:
            return pd.DataFrame()
        
        try:
            logger.info(f"Processing image with EasyOCR: {jpeg_path}")
            
            # Initialize EasyOCR reader (first time will download models)
            if not hasattr(self, 'reader'):
                logger.info("Initializing EasyOCR reader (this may take a moment on first run)...")
                self.reader = easyocr.Reader(['es', 'en'], gpu=False)
            
            # Perform OCR
            result = self.reader.readtext(jpeg_path, detail=0)  # detail=0 returns just text
            
            # Join all text
            text = '\n'.join(result)
            logger.info(f"OCR extracted {len(text)} characters from {jpeg_path}")
            
            # Parse text to extract structured data
            return self._parse_ocr_text_structured(text, jpeg_path)
            
        except Exception as e:
            logger.error(f"Error in _extract_from_jpeg: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return pd.DataFrame()
    
    def _parse_ocr_text_structured(self, text, source_file):
        """
        Parse OCR text from Zaloze price list images.
        The images show tables with columns: Diametro, Product Types, Prices
        """
        data = []
        lines = text.split('\n')
        
        # Look for table data patterns
        # Format: size followed by prices for different product types
        current_size = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Try to identify size (e.g., 1/2, 3/4, 1, 1 1/4, 2, etc.)
            size_match = re.match(r'^(\d+(?:\s*\d*/\d+)?|\d+/\d+)$', line.replace('½', '1/2').replace('¾', '3/4'))
            if size_match:
                current_size = size_match.group(1).strip()
                continue
            
            # Try to extract prices (numbers with decimals)
            prices = re.findall(r'\d+\.\d{2}', line)
            
            if current_size and prices:
                # Map prices to product types based on column position
                # From images: Codo Radio Largo 90° (Std/Ex Pesado), Codo 45° (Std/Ex Pesado), etc.
                
                for j, price in enumerate(prices):
                    try:
                        precio_float = float(price)
                        
                        # Determine product type based on column index
                        if j == 0:
                            tipo = 'CODO RADIO LARGO 90°'
                            espesor = 'STD'
                        elif j == 1:
                            tipo = 'CODO RADIO LARGO 90°'
                            espesor = 'XS'
                        elif j == 2:
                            tipo = 'CODO 45°'
                            espesor = 'STD'
                        elif j == 3:
                            tipo = 'CODO 45°'
                            espesor = 'XS'
                        elif j == 4:
                            tipo = 'CODO RADIO CORTO 90°'
                            espesor = 'STD'
                        elif j == 5:
                            tipo = 'CODO RADIO CORTO 90°'
                            espesor = 'XS'
                        else:
                            continue
                        
                        data.append({
                            'codigo': f"ZALOZE-{tipo}-{current_size}-{espesor}",
                            'descripcion': f"{tipo} {current_size} {espesor}",
                            'precio': precio_float,
                            'tipo_serie': tipo,
                            'espesor': espesor,
                            'size': current_size
                        })
                    except ValueError:
                        continue
        
        if not data:
            # If structured parsing failed, try simple pattern matching
            logger.warning("Structured parsing failed, attempting simple pattern matching")
            return self._parse_ocr_text(text)
        
        df = pd.DataFrame(data)
        logger.info(f"Extracted {len(df)} products from OCR of {source_file}")
        return df
    
    def _parse_ocr_text(self, text):
        """Parse OCR text to extract product data."""
        lines = text.split('\n')
        data = []
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 10:
                continue
            
            # Look for patterns: code, description, price
            match = re.search(r'([A-Z0-9\-]+)\s+(.+?)\s+\$?\s*([\d,.]+)', line)
            if match:
                data.append({
                    'codigo': match.group(1),
                    'descripcion': match.group(2).strip(),
                    'precio': match.group(3).replace(',', '.')
                })
        
        if not data:
            return pd.DataFrame()
        
        df = pd.DataFrame(data)
        return self._normalize_zaloze_data(df)
    
    def _normalize_zaloze_data(self, df):
        """Normalize Zaloze data to standard format."""
        result_df = pd.DataFrame()
        
        # Find codigo column
        codigo_col = self._find_column(df, ['codigo', 'code', 'cod', 'item'])
        if codigo_col:
            result_df['codigo'] = df[codigo_col].astype(str)
        elif len(df.columns) > 0:
            result_df['codigo'] = df.iloc[:, 0].astype(str)
        
        # Find descripcion column
        desc_col = self._find_column(df, ['descripcion', 'description', 'desc'])
        if desc_col:
            result_df['descripcion'] = df[desc_col].astype(str)
        elif len(df.columns) > 1:
            result_df['descripcion'] = df.iloc[:, 1].astype(str)
        
        # Find precio column
        precio_col = self._find_column(df, ['precio', 'price', 'valor'])
        if precio_col:
            result_df['precio'] = pd.to_numeric(df[precio_col].astype(str).str.replace(',', '.'), 
                                               errors='coerce')
        elif len(df.columns) > 2:
            result_df['precio'] = pd.to_numeric(df.iloc[:, 2].astype(str).str.replace(',', '.'), 
                                               errors='coerce')
        
        # Extract tipo_serie, espesor, size
        if 'codigo' in result_df.columns:
            result_df['tipo_serie'] = result_df['codigo'].apply(self._extract_tipo_serie)
        else:
            result_df['tipo_serie'] = ''
        
        if 'descripcion' in result_df.columns:
            result_df['espesor'] = result_df['descripcion'].apply(self._extract_espesor)
            result_df['size'] = result_df['descripcion'].apply(self._extract_size)
        else:
            result_df['espesor'] = ''
            result_df['size'] = ''
        
        # Remove invalid rows
        result_df = result_df.dropna(subset=['precio'])
        result_df = result_df[result_df['precio'] > 0]
        
        return result_df
    
    def _find_column(self, df, possible_names):
        """Find column by possible names."""
        df_cols = [str(c).strip().lower() for c in df.columns]
        for name in possible_names:
            if name in df_cols:
                idx = df_cols.index(name)
                return df.columns[idx]
        return None
    
    def _extract_tipo_serie(self, codigo):
        """Extract tipo_serie from product code."""
        try:
            parts = str(codigo).split('-')
            if len(parts) > 0:
                return parts[0].strip().upper()
        except:
            pass
        return ''
    
    def _extract_espesor(self, description):
        """Extract thickness from description."""
        try:
            desc = str(description).upper()
            match = re.search(r'SCH\s*(\d+|STD|XS|XXS)', desc)
            if match:
                return match.group(0).replace(' ', '')
        except:
            pass
        return ''
    
    def _extract_size(self, description):
        """Extract size from description."""
        try:
            desc = str(description)
            match = re.search(r'(\d+/?\d*)\s*["\']', desc)
            if match:
                return match.group(1)
            match = re.search(r'(\d+)\s*(MM|INCH)', desc, re.IGNORECASE)
            if match:
                return match.group(1) + match.group(2).upper()
        except:
            pass
        return ''

